import os
from flask import Flask, render_template, request, redirect, url_for, flash, session, send_from_directory
from flask_sqlalchemy import SQLAlchemy
from flask_session import Session
import gspread
from oauth2client.service_account import ServiceAccountCredentials # Para gspread com service account
import pandas as pd # Para facilitar a manipulação dos dados da planilha
from google.oauth2 import service_account # Para API do Drive
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import io
import logging
from docx import Document
from docx.shared import Pt # Se precisar manipular tamanhos de fonte, etc.
from docx2pdf import convert as convert_docx_to_pdf
import datetime # Para datas
from num2words import num2words
import subprocess


# Configuração inicial do Flask
app = Flask(__name__)
app.logger.setLevel(logging.DEBUG) # Ou logging.INFO

# Configurações (podem ir para um arquivo de configuração ou .env)
# Pasta para uploads
UPLOAD_FOLDER = 'uploads'
CREDENTIALS_FILE = './utils/credentials-google.json' # Caminho para o arquivo de credenciais

# Chave secreta para sessões e mensagens flash (importante para produção)
app.config['SECRET_KEY'] = 'BusquemConhecimento2025%$BdmDigital' # Mude isso para uma chave segura!
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.template_folder = 'templates'
app.static_folder = 'static'

# Configurações do SQLAlchemy (usando SQLite)
basedir = os.path.abspath(os.path.dirname(__file__)) # Diretório base da aplicação
# Define o caminho para o arquivo do banco de dados SQLite
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'flask_session.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False # Desativa warnings desnecessários

# Configurações do Flask-Session
app.config['SESSION_TYPE'] = 'sqlalchemy' # Indica que usaremos SQLAlchemy para armazenar as sessões
app.config['SESSION_PERMANENT'] = False # Sessões expiram quando o navegador é fechado (ou defina True e SESSION_LIFETIME)
app.config['SESSION_USE_SIGNER'] = True  # Assina o cookie da sessão para segurança
app.config['SESSION_SQLALCHEMY_TABLE'] = 'sessions' # Nome da tabela no DB para armazenar as sessões

# Inicializa o SQLAlchemy
db = SQLAlchemy(app)

# Diz ao Flask-Session para usar sua instância 'db' do SQLAlchemy
app.config['SESSION_SQLALCHEMY'] = db

# Inicializa a extensão Flask-Session DEPOIS de configurar SESSION_TYPE e outras configs
server_session = Session(app) # Isso substitui a gestão de sessão padrão do Flask

# --- Fim da Configuração para Flask-Session e SQLAlchemy ---

# Escopos para as APIs do Google
SCOPES_GSPREAD = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
SCOPES_DRIVE = ['https://www.googleapis.com/auth/drive.readonly'] # ou drive completo se precisar criar/modificar

# --- Criação da Tabela de Sessões ---
# Isso garante que a tabela 'sessions' exista no banco de dados SQLite.
# db.create_all() deve ser chamado dentro de um contexto de aplicação.
with app.app_context():
    db.create_all()
    app.logger.info("Tabela de sessões verificada/criada no banco de dados SQLite.")


def get_google_sheets_client():
    """Autentica com a API do Google Sheets usando gspread e credenciais de serviço."""
    try:
        if not os.path.exists(CREDENTIALS_FILE):
            app.logger.error(f"Arquivo de credenciais '{CREDENTIALS_FILE}' não encontrado.")
            flash(f"ERRO INTERNO: Arquivo de credenciais '{CREDENTIALS_FILE}' não encontrado. Contate o administrador.", "danger")
            return None
        
        creds = ServiceAccountCredentials.from_json_keyfile_name(CREDENTIALS_FILE, SCOPES_GSPREAD)
        client = gspread.authorize(creds)
        app.logger.info("Cliente Google Sheets autenticado com sucesso.")
        return client
    except Exception as e:
        app.logger.error(f"Erro ao autenticar com Google Sheets: {e}", exc_info=True)
        flash(f"ERRO INTERNO: Falha na autenticação com Google Sheets. Verifique os logs. Detalhe: {e}", "danger")
        return None

def get_google_drive_service():
    """Autentica com a API do Google Drive."""
    try:
        if not os.path.exists(CREDENTIALS_FILE):
            app.logger.error(f"Arquivo de credenciais '{CREDENTIALS_FILE}' não encontrado.")
            flash(f"ERRO INTERNO: Arquivo de credenciais '{CREDENTIALS_FILE}' não encontrado. Contate o administrador.", "danger")
            return None

        creds = service_account.Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES_DRIVE)
        service = build('drive', 'v3', credentials=creds)
        app.logger.info("Serviço Google Drive autenticado com sucesso.")
        return service
    except Exception as e:
        app.logger.error(f"Erro ao autenticar com Google Drive: {e}", exc_info=True)
        flash(f"ERRO INTERNO: Falha na autenticação com Google Drive. Verifique os logs. Detalhe: {e}", "danger")
        return None

def get_sheet_data(sheet_url_or_id):
    """Busca dados de uma Planilha Google e retorna como DataFrame."""
    app.logger.info(f"Tentando obter cliente Google Sheets para: {sheet_url_or_id}")
    client = get_google_sheets_client()
    if not client:
        # A função get_google_sheets_client() já deve ter emitido um flash e logado o erro.
        return None
    
    try:
        app.logger.info(f"Tentando abrir planilha com identificador: '{sheet_url_or_id}'")
        # gspread.open_by_url() ou gspread.open_by_key() são mais explícitos
        # Se o usuário sempre fornecer o ID, open_by_key é melhor.
        # Se for URL, open_by_url. gspread.open() tenta adivinhar.
        
        # Vamos assumir que o usuário está fornecendo o ID diretamente
        # Se for uma URL completa, você pode precisar extrair o ID dela primeiro.
        # Ex: "https://docs.google.com/spreadsheets/d/ESTE_EH_O_ID/edit#gid=0" -> ID: "ESTE_EH_O_ID"
        
        # Simplificação: Se for uma URL, tente abrir por URL, senão, por chave (ID)
        actual_id_to_open = sheet_url_or_id
        if "docs.google.com/spreadsheets/d/" in sheet_url_or_id:
            actual_id_to_open = sheet_url_or_id.split('/d/')[1].split('/')[0]
            app.logger.info(f"ID extraído da URL: {actual_id_to_open}")
            sheet = client.open_by_key(actual_id_to_open) # Usar open_by_key com o ID extraído
        else: # Assume que é um ID direto
            app.logger.info(f"Assumindo que '{sheet_url_or_id}' é um ID de planilha.")
            sheet = client.open_by_key(sheet_url_or_id)

        app.logger.info(f"Planilha '{sheet.title}' aberta com sucesso.")
        
        worksheet = sheet.sheet1 # Assume a primeira aba
        app.logger.info(f"Acessando primeira aba (worksheet): '{worksheet.title}'.")
        
        data = worksheet.get_all_records() # Espera cabeçalhos na primeira linha
        app.logger.info(f"Dados lidos da planilha: {len(data)} registros.")
        
        if not data:
            app.logger.warning(f"Nenhum dado encontrado na planilha '{sheet.title}' (worksheet: '{worksheet.title}'). Verifique se há dados e cabeçalhos.")
            # Não é necessariamente um erro que impede o fluxo, mas um aviso.
            # flash("Atenção: Nenhum dado encontrado na planilha ou ela está vazia.", "warning")

        return pd.DataFrame(data)

    except gspread.exceptions.SpreadsheetNotFound:
        app.logger.error(f"SpreadsheetNotFound para o identificador: '{sheet_url_or_id}' (ID tentado: '{actual_id_to_open if 'actual_id_to_open' in locals() else sheet_url_or_id}').")
        flash(f"ERRO: Planilha não encontrada com o identificador fornecido. Verifique a URL/ID e se a planilha foi compartilhada com o e-mail da conta de serviço ({CREDENTIALS_FILE}).", "danger")
        return None
    except gspread.exceptions.APIError as api_e:
        app.logger.error(f"gspread.exceptions.APIError ao acessar planilha '{sheet_url_or_id}': {api_e}", exc_info=True)
        flash(f"ERRO DE API ao acessar a planilha: {api_e}. Verifique as permissões de compartilhamento ou problemas na API do Google.", "danger")
        return None
    except Exception as e:
        app.logger.error(f"Erro inesperado ao ler planilha '{sheet_url_or_id}': {e}", exc_info=True)
        flash(f"ERRO INESPERADO ao ler a planilha: {e}. Verifique os logs do servidor.", "danger")
        return None

def download_google_doc_as_docx(doc_url_or_id, output_folder):
    """Baixa um Google Doc como arquivo .docx e o salva na output_folder."""
    app.logger.info(f"Tentando obter serviço Google Drive para baixar doc: {doc_url_or_id}")
    service = get_google_drive_service()
    if not service:
        # A função get_google_drive_service() já deve ter emitido um flash e logado o erro.
        return None

    file_id = doc_url_or_id
    original_input_id = doc_url_or_id # Para logging

    # Extrai o File ID da URL se for uma URL completa
    if "docs.google.com/document/d/" in doc_url_or_id:
        try:
            file_id = doc_url_or_id.split('/d/')[1].split('/')[0]
            app.logger.info(f"ID '{file_id}' extraído da URL do documento: '{doc_url_or_id}'")
        except IndexError:
            app.logger.error(f"Não foi possível extrair o ID da URL do documento: '{doc_url_or_id}'")
            flash(f"URL do Documento Google inválida: '{doc_url_or_id}'. Não foi possível extrair o ID.", "danger")
            return None
    else:
        app.logger.info(f"Assumindo que '{doc_url_or_id}' é um ID de documento direto.")

    try:
        app.logger.info(f"Solicitando exportação do Google Doc ID '{file_id}' como DOCX.")
        request_body = service.files().export_media(
            fileId=file_id,
            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Obter o nome original do arquivo para usar no nome do arquivo salvo
        file_metadata = service.files().get(fileId=file_id, fields='name').execute()
        original_filename = file_metadata.get('name', 'documento_contrato_sem_nome')
        # Limpa o nome do arquivo para evitar caracteres problemáticos e adiciona .docx
        safe_filename_base = "".join([c if c.isalnum() or c in (' ', '_', '-') else "_" for c in original_filename])
        safe_filename = f"{safe_filename_base}.docx"
        
        filepath = os.path.join(output_folder, safe_filename)

        app.logger.info(f"Iniciando download do DOCX para: '{filepath}'")
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request_body)
        done = False
        while not done:
            status, done = downloader.next_chunk()
            if status:
                app.logger.info(f"Download do DOCX {int(status.progress() * 100)}% concluído.")
        
        fh.seek(0) # Voltar para o início do buffer de bytes
        with open(filepath, 'wb') as f:
            f.write(fh.read())
        
        app.logger.info(f"Documento '{original_filename}' (ID: {file_id}) baixado com sucesso como '{filepath}'.")
        return filepath

    except Exception as e:
        # Verifica se o erro é específico de "File not found" ou permissão
        error_details = str(e)
        if "File not found" in error_details or "notFound" in error_details:
            app.logger.error(f"Google Doc com ID '{file_id}' não encontrado (originalmente '{original_input_id}'). Erro: {e}", exc_info=True)
            flash(f"ERRO: Documento Google com ID '{file_id}' não encontrado. Verifique o ID/URL e o compartilhamento.", "danger")
        elif "insufficient permissions" in error_details.lower() or "forbidden" in error_details.lower():
            app.logger.error(f"Permissão insuficiente para acessar o Google Doc ID '{file_id}'. Erro: {e}", exc_info=True)
            flash(f"ERRO: Permissão insuficiente para acessar o Documento Google (ID: {file_id}). Verifique o compartilhamento.", "danger")
        else:
            app.logger.error(f"Erro ao baixar Google Doc '{original_input_id}' (ID tentado: {file_id}): {e}", exc_info=True)
            flash(f"ERRO INESPERADO ao baixar o Documento Google: {e}. Verifique os logs.", "danger")
        return None


@app.route('/', methods=['GET', 'POST'])
def index():
    app.logger.info("--- FUNÇÃO INDEX ACESSADA ---")
    donatarios_df = None
    path_template_docx = None
    donatarios_para_exibicao = [] # Inicializa como lista vazia

    sheet_url_value = session.get('latest_sheet_url', '')
    doc_url_value = session.get('latest_doc_url', '')

    if request.method == 'POST':
        sheet_url_from_form = request.form.get('sheet_url')
        doc_url_from_form = request.form.get('doc_url')
        sheet_url_value = sheet_url_from_form
        doc_url_value = doc_url_from_form
        session['latest_sheet_url'] = sheet_url_from_form
        session['latest_doc_url'] = doc_url_from_form

        if not sheet_url_from_form or not doc_url_from_form:
            flash("Por favor, forneça as URLs/IDs da Planilha Google e do Documento Google.", "danger")
            session.pop('donatarios_data_json', None)
            session.pop('latest_doc_template_path', None)
        else:
            app.logger.info(f"Carregando dados da planilha: {sheet_url_from_form}")
            donatarios_df = get_sheet_data(sheet_url_from_form)

            if donatarios_df is not None:
                app.logger.info(f"donatarios_df carregado com {len(donatarios_df)} linhas.")
                # ATENÇÃO: O AVISO DE COOKIE GRANDE AINDA OCORRERÁ AQUI!
                session['donatarios_data_json'] = donatarios_df.to_json(orient='records')
                app.logger.info("Dados completos dos donatários salvos na sessão (pode exceder o limite do cookie).")
            else:
                app.logger.warning("Falha ao carregar donatarios_df da planilha.")
                session.pop('donatarios_data_json', None)

            app.logger.info(f"Baixando template do documento: {doc_url_from_form}")
            path_template_docx = download_google_doc_as_docx(doc_url_from_form, app.config['UPLOAD_FOLDER'])

            if path_template_docx is not None:
                app.logger.info(f"path_template_docx baixado: {path_template_docx}")
                session['latest_doc_template_path'] = path_template_docx
            else:
                app.logger.warning("Falha ao baixar path_template_docx.")
                session.pop('latest_doc_template_path', None)

            if donatarios_df is not None and path_template_docx is not None:
                flash("Planilha e Documento carregados com sucesso!", "success")
            elif donatarios_df is not None and path_template_docx is None:
                flash("Planilha carregada, mas houve erro ao carregar o Documento.", "warning")
            elif donatarios_df is None and path_template_docx is not None:
                flash("Documento carregado, mas houve erro ao carregar a Planilha.", "warning")
    
    # --- Lógica Comum para GET e para continuar após POST ---
    # Tenta carregar donatarios_df da sessão se não foi carregado no POST atual
    if donatarios_df is None and 'donatarios_data_json' in session:
        app.logger.info("Tentando carregar donatarios_df da sessão (GET ou POST falhou em carregar dados novos).")
        try:
            # ATENÇÃO: SE A SESSÃO FOI TRUNCADA/CORROMPIDA, ISTO PODE FALHAR OU RETORNAR DADOS INCOMPLETOS
            donatarios_df = pd.read_json(session['donatarios_data_json'], orient='records')
            app.logger.info(f"donatarios_df carregado da sessão com {len(donatarios_df)} linhas.")
        except Exception as e:
            app.logger.error(f"Erro ao ler 'donatarios_data_json' da sessão (possivelmente corrompido): {e}")
            session.pop('donatarios_data_json', None)
            donatarios_df = None

    if path_template_docx is None and 'latest_doc_template_path' in session:
        candidate_path = session['latest_doc_template_path']
        if os.path.exists(candidate_path):
            path_template_docx = candidate_path
            app.logger.info(f"path_template_docx carregado da sessão: {path_template_docx}")
        else:
            session.pop('latest_doc_template_path', None)

    # Processa donatarios_df para exibição (SEMPRE que donatarios_df tiver dados)
    if donatarios_df is not None and not donatarios_df.empty:
        app.logger.info("Processando donatarios_df para criar donatarios_para_exibicao.")
        
        nome_coluna_para_nome_do_donatario = 'NOME' # Confirmado pelo usuário
        nome_coluna_para_cpf_do_donatario = 'CPF'   # Confirmado pelo usuário

        colunas_identificadas_para_exibicao = []
        if nome_coluna_para_nome_do_donatario in donatarios_df.columns:
            colunas_identificadas_para_exibicao.append(nome_coluna_para_nome_do_donatario)
        else:
            app.logger.warning(f"ALERTA: Coluna '{nome_coluna_para_nome_do_donatario}' NÃO encontrada na planilha.")
        
        if nome_coluna_para_cpf_do_donatario in donatarios_df.columns:
            colunas_identificadas_para_exibicao.append(nome_coluna_para_cpf_do_donatario)
        else:
            app.logger.warning(f"ALERTA: Coluna '{nome_coluna_para_cpf_do_donatario}' NÃO encontrada na planilha.")
        
        if colunas_identificadas_para_exibicao:
            donatarios_para_exibicao = donatarios_df[colunas_identificadas_para_exibicao].to_dict(orient='records')
            app.logger.info(f"donatarios_para_exibicao preparado com {len(donatarios_para_exibicao)} registros e colunas: {list(donatarios_para_exibicao[0].keys()) if donatarios_para_exibicao else 'VAZIO'}.")
        else:
            app.logger.error("ERRO: Nenhuma das colunas (NOME/CPF) foi encontrada. Tabela de donatarios ficará vazia.")
            donatarios_para_exibicao = [] # Garante que é uma lista vazia
    else:
        app.logger.info("donatarios_df está vazio ou None ao final do processamento. 'donatarios_para_exibicao' será uma lista vazia.")
        donatarios_para_exibicao = [] # Garante que é uma lista vazia se não houver dados

    # LOG FINAL ANTES DE RENDERIZAR:
    app.logger.info(f"Antes de renderizar: len(donatarios_para_exibicao) = {len(donatarios_para_exibicao)}")
    if donatarios_para_exibicao:
         app.logger.info(f"Primeiro item em donatarios_para_exibicao: {donatarios_para_exibicao[0]}")


    return render_template('index.html', 
                           donatarios_exibicao=donatarios_para_exibicao,
                           path_template_docx=path_template_docx,
                           sheet_url_value=sheet_url_value,
                           doc_url_value=doc_url_value)


@app.route('/gerar_contrato', methods=['POST'])
def gerar_contrato():
    app.logger.info("--- ROTA GERAR_CONTRATO ACESSADA ---")

    # 1. Obter dados do formulário
    selected_donatario_index_str = request.form.get('donatario_selecionado_index')
    valor_doacao_str = request.form.get('valor_doacao')
    aliquota_str = request.form.get('aliquota') # Alíquota em %, ex: usuário digita 4 para 4%
    doc_template_path = request.form.get('doc_template_path')

    app.logger.info(f"Formulário recebido - Índice Str: {selected_donatario_index_str}, Valor Str: {valor_doacao_str}, Alíquota Str: {aliquota_str}, Template: {doc_template_path}")

    # 2. Validação e conversão dos dados do formulário
    if not all([selected_donatario_index_str, valor_doacao_str, aliquota_str, doc_template_path]):
        flash("Dados incompletos recebidos para gerar o contrato. Tente novamente.", "danger")
        app.logger.warning("Dados incompletos no formulário para gerar_contrato.")
        return redirect(url_for('index'))

    try:
        selected_donatario_index = int(selected_donatario_index_str)
        valor_bruto_doacao = float(valor_doacao_str)
        aliquota_percentual = float(aliquota_str) # Ex: 4.0
        app.logger.info(f"Valores convertidos - Índice: {selected_donatario_index}, Valor Bruto: {valor_bruto_doacao}, Alíquota: {aliquota_percentual}%")
    except ValueError:
        flash("Valores inválidos para índice, doação ou alíquota. Use números válidos.", "danger")
        app.logger.error(f"ValueError ao converter dados: índice='{selected_donatario_index_str}', valor='{valor_doacao_str}', alíquota='{aliquota_str}'", exc_info=True)
        return redirect(url_for('index'))
    except TypeError:
        flash("Erro nos tipos de dados recebidos (provavelmente algum valor está faltando).", "danger")
        app.logger.error(f"TypeError nos dados do formulário: índice='{selected_donatario_index_str}', valor='{valor_doacao_str}', alíquota='{aliquota_str}'", exc_info=True)
        return redirect(url_for('index'))

    # 3. Recuperar dados completos do donatário da sessão
    donatarios_json_str = session.get('donatarios_data_json')
    if not donatarios_json_str:
        flash("Sessão expirada ou dados dos donatários não encontrados. Por favor, recarregue a planilha.", "warning")
        app.logger.warning("donatarios_data_json não encontrado na sessão.")
        return redirect(url_for('index'))

    try:
        all_donatarios_list = pd.read_json(donatarios_json_str, orient='records').to_dict(orient='records')
    except Exception as e:
        flash("Erro ao processar dados dos donatários da sessão. Recarregue a planilha.", "danger")
        app.logger.error(f"Erro ao fazer pd.read_json de donatarios_data_json: {e}", exc_info=True)
        return redirect(url_for('index'))

    if not (0 <= selected_donatario_index < len(all_donatarios_list)):
        flash(f"Índice de donatário selecionado ({selected_donatario_index}) inválido.", "danger")
        app.logger.error(f"Índice de donatário selecionado inválido: {selected_donatario_index}. Lista com {len(all_donatarios_list)} itens.")
        return redirect(url_for('index'))

    selected_donatario_data = all_donatarios_list[selected_donatario_index]
    app.logger.info(f"Dados Completos do Donatário Selecionado: {selected_donatario_data}")

    # 4. Processar o template DOCX
    app.logger.info(f"Iniciando processamento do template DOCX: {doc_template_path}")
    try:
        document = Document(doc_template_path)

        # Cálculos
        valor_itcmd = (valor_bruto_doacao * aliquota_percentual) / 100.0
        valor_liquido_doacao = valor_bruto_doacao - valor_itcmd

        # Data por extenso (Exemplo: Mossoró/RN, 12 de maio de 2025)
        # Você precisará definir a CIDADE_DOADOR e ESTADO_DOADOR, talvez da planilha ou fixo.
        # Por enquanto, vou usar um valor fixo como no seu exemplo.
        # Adapte conforme sua necessidade para pegar a cidade/estado do doador dinamicamente se preciso.
        cidade_doador = "Mossoró" # Ou buscar de algum lugar
        uf_doador = "RN"          # Ou buscar de algum lugar
        data_atual = datetime.date.today()
        meses_extenso = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", 
                         "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
        data_formatada_extenso = f"{cidade_doador}/{uf_doador}, {data_atual.day} de {meses_extenso[data_atual.month - 1]} de {data_atual.year}"

        # Valores por extenso
        # O 'to='currency'' adiciona "reais" e "centavos"
        # Se quiser só o número por extenso, use to='cardinal' ou to='spellout'
        # e adicione "reais" e "centavos" manualmente se necessário.
        try:
            valor_bruto_extenso = num2words(valor_bruto_doacao, lang='pt_BR', to='currency')
            valor_itcmd_extenso = num2words(valor_itcmd, lang='pt_BR', to='currency')
            valor_liquido_extenso = num2words(valor_liquido_doacao, lang='pt_BR', to='currency')
        except Exception as e_num2words:
            app.logger.error(f"Erro ao converter números para extenso com num2words: {e_num2words}", exc_info=True)
            flash(f"Erro ao gerar valores por extenso: {e_num2words}. Verifique os valores numéricos.", "danger")
            # Você pode optar por continuar sem os valores por extenso ou parar.
            # Por agora, vamos definir como "ERRO NA GERAÇÃO" para que o processo continue para depuração.
            valor_bruto_extenso = valor_itcmd_extenso = valor_liquido_extenso = "ERRO NA GERAÇÃO DO VALOR POR EXTENSO"


        # --- Dicionário de Substituições ---
        # !! ADAPTE AS CHAVES DE selected_donatario_data PARA CORRESPONDER AOS NOMES DAS SUAS COLUNAS NA PLANILHA !!
        substituicoes = {
            "<<NOME_DONATARIO>>": str(selected_donatario_data.get('NOME', '')).strip(),
            "<<NACIONALIDADE_DONATARIO>>": str(selected_donatario_data.get('NACIONALIDADE', 'N/D')).lower().strip(), # Exemplo: BRASILEIRA
            "<<ESTADO_CIVIL_DONATARIO>>": str(selected_donatario_data.get('ESTADO_CIVIL', 'N/D')).lower().strip(),
            "<<PROFISSAO_DONATARIO>>": str(selected_donatario_data.get('PROFISSAO', 'N/D')).lower().strip(),
            "<<RG_DONATARIO>>": str(selected_donatario_data.get('RG', 'N/D')).strip(),
            "<<CPF_DONATARIO>>": str(selected_donatario_data.get('CPF', '')).strip(),
            "<<ENDERECO_DONATARIO>>": str(selected_donatario_data.get('ENDERECO', 'N/D')).strip(), # Ex: "RUA 4, CHÁCARA 300, LOTE, 3-B"
            "<<CIDADE_UF_DONATARIO>>": str(selected_donatario_data.get('CIDADE_UF', 'N/D')).strip(), # Ex: "BRASILIA/DF"
            "<<CEP_DONATARIO>>": str(selected_donatario_data.get('CEP', 'N/D')).strip(),
            "<<TELEFONE_DONATARIO>>": str(selected_donatario_data.get('TELEFONE', 'N/D')).strip(),
            "<<EMAIL_DONATARIO>>": str(selected_donatario_data.get('EMAIL', 'N/D')).upper().strip(),
            "<<BANCO_DONATARIO>>": str(selected_donatario_data.get('BANCO', 'N/D')).strip(),
            "<<AGENCIA_DONATARIO>>": str(selected_donatario_data.get('AGENCIA', 'N/D')).strip(),
            "<<CONTA_DONATARIO>>": str(selected_donatario_data.get('CONTA', 'N/D')).strip(),
            "<<CONTA_TIPO>>": str(selected_donatario_data.get('OPERACAO', 'N/D')).lower().strip(),
            
            "<<VALOR_BRUTO_DOACAO_NUM>>": f"{valor_bruto_doacao:,.2f}".replace('.', '#').replace(',', '.').replace('#', ','), # Formato 1.234,56
            "<<VALOR_BRUTO_DOACAO_EXTENSO>>": valor_bruto_extenso.lower(),
            "<<ALIQUOTA_ITCMD_PERCENTUAL>>": f"{aliquota_percentual:,.2f}%".replace('.', ','), # Formato 4,00%
            "<<VALOR_ITCMD_NUM>>": f"{valor_itcmd:,.2f}".replace('.', '#').replace(',', '.').replace('#', ','),
            "<<VALOR_ITCMD_EXTENSO>>": valor_itcmd_extenso.lower(),
            "<<VALOR_LIQUIDO_DOACAO_NUM>>": f"{valor_liquido_doacao:,.2f}".replace('.', '#').replace(',', '.').replace('#', ','),
            "<<VALOR_LIQUIDO_DOACAO_EXTENSO>>": valor_liquido_extenso.lower(),
            
            "<<LOCAL_DATA_COMPLETA>>": data_formatada_extenso,
        }
        
        app.logger.info(f"Dicionário de substituições montado: {substituicoes}")

        # Substituição em parágrafos
        for paragraph in document.paragraphs:
            for run in paragraph.runs: # Iterar sobre 'runs' para melhor preservação da formatação
                text = run.text
                for key, value in substituicoes.items():
                    if key in text:
                        text = text.replace(key, value)
                run.text = text # Atualiza o texto do 'run'

        # Substituição em tabelas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            for key, value in substituicoes.items():
                                if key in text:
                                    text = text.replace(key, value)
                            run.text = text
        
        # 5. Salvar o documento DOCX modificado
        nome_donatario_arquivo = "".join(c if c.isalnum() else "_" for c in selected_donatario_data.get('NOME', 'donatario_sem_nome'))
        nome_arquivo_final = f"CONTRATO_{nome_donatario_arquivo}_{data_atual.strftime('%Y%m%d')}.docx"
        
        pasta_contratos_gerados = os.path.join(app.config['UPLOAD_FOLDER'], 'contratos_gerados')
        if not os.path.exists(pasta_contratos_gerados):
            os.makedirs(pasta_contratos_gerados)
            app.logger.info(f"Pasta '{pasta_contratos_gerados}' criada.")

        caminho_docx_final = os.path.join(pasta_contratos_gerados, nome_arquivo_final)
        
        document.save(caminho_docx_final)
        app.logger.info(f"Contrato DOCX modificado salvo em: {caminho_docx_final}")

        # --- ETAPA 5: Converter DOCX para PDF ---
        nome_arquivo_pdf = nome_arquivo_final.replace(".docx", ".pdf")
        caminho_pdf_final = os.path.join(pasta_contratos_gerados, nome_arquivo_pdf)
        
        # Diretório onde o PDF será salvo (para o comando do LibreOffice)
        output_dir_for_pdf = pasta_contratos_gerados

        try:
            app.logger.info(f"Iniciando conversão de '{caminho_docx_final}' para PDF em '{output_dir_for_pdf}' usando SEU COMANDO ATUAL...")
            
            # !!! USE O SEU COMANDO ATUAL AQUI, O QUE ESTÁ GERANDO O PDF !!!
            comando = [
                "libreoffice", # ou "soffice", ou o que você tem que está funcionando
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir_for_pdf,
                caminho_docx_final
            ]
            # Se você usa xvfb-run, seria:
            # comando = [
            #     "xvfb-run", "-a",
            #     "libreoffice",
            #     "--headless",
            #     "--convert-to", "pdf",
            #     "--outdir", output_dir_for_pdf,
            #     caminho_docx_final
            # ]
            
            app.logger.info(f"Executando comando: {' '.join(comando)}")
            process = subprocess.run(comando, timeout=60, capture_output=True, text=True, check=False)

            app.logger.info(f"Comando de conversão finalizado. Código de retorno: {process.returncode}")
            app.logger.info(f"Stdout do comando: {process.stdout}")
            app.logger.info(f"Stderr do comando: {process.stderr}")

            if process.returncode == 0 and os.path.exists(caminho_pdf_final):
                app.logger.info(f"Arquivo PDF gerado com sucesso e encontrado: {caminho_pdf_final}")
                flash(f"Contrato PDF para {selected_donatario_data.get('NOME')} gerado com sucesso!", "success")
                return render_template('sucesso_geracao.html', 
                                       nome_arquivo_pdf=nome_arquivo_pdf,
                                       nome_donatario=selected_donatario_data.get('NOME'))
            else:
                # O PDF foi gerado, mas algo na condição falhou, ou o returncode não foi 0
                app.logger.error(f"Falha na condição de sucesso pós-conversão. Código de retorno: {process.returncode}. PDF existe? {os.path.exists(caminho_pdf_final)}")
                if not os.path.exists(caminho_pdf_final) and process.returncode == 0:
                     app.logger.error(f"Comando de conversão retornou 0, mas o arquivo PDF não foi encontrado em '{caminho_pdf_final}'.")
                
                mensagem_erro_flash = "Contrato DOCX gerado, mas houve um problema na finalização da conversão para PDF."
                if process.stderr:
                    mensagem_erro_flash += f" Detalhe do conversor: {process.stderr[:100]}" # Primeiros 100 caracteres do erro
                elif process.returncode != 0:
                    mensagem_erro_flash += f" Código de erro do conversor: {process.returncode}."

                flash(mensagem_erro_flash, "danger")
                return redirect(url_for('index'))

        except FileNotFoundError as e_fnf:
            app.logger.error(f"Comando para conversão PDF não encontrado (ex: libreoffice): {e_fnf}", exc_info=True)
            flash(f"Erro ao converter para PDF: Ferramenta de conversão ('{comando[0]}') não encontrada. {e_fnf}", "danger")
            return redirect(url_for('index'))
        except subprocess.TimeoutExpired as e_timeout:
            app.logger.error("Timeout ao converter DOCX para PDF.", exc_info=True)
            flash("A conversão para PDF demorou demais (timeout). Tente novamente.", "danger")
            return redirect(url_for('index'))
        except Exception as e_pdf:
            app.logger.error(f"CAPTURA BRUTA - Exceção durante ou após conversão para PDF: {type(e_pdf).__name__} - {str(e_pdf)}", exc_info=True)
            
            mensagem_erro_flash = f"Erro inesperado durante conversão para PDF: {type(e_pdf).__name__}."
            if str(e_pdf):
                 mensagem_erro_flash += f" Detalhe: {str(e_pdf)[:100]}"

            flash(f"Contrato DOCX gerado ({nome_arquivo_final}), mas ocorreu um erro inesperado. {mensagem_erro_flash}", "danger")
            return redirect(url_for('index'))


    except FileNotFoundError:
        app.logger.error(f"Erro: Template DOCX não encontrado em '{doc_template_path}'", exc_info=True)
        flash(f"Erro crítico: O arquivo de template DOCX não foi encontrado. Verifique o download.", "danger")
        return redirect(url_for('index'))
    except Exception as e:
        app.logger.error(f"Erro ao processar o contrato DOCX: {e}", exc_info=True)
        flash(f"Ocorreu um erro inesperado ao gerar o contrato: {e}", "danger")
        return redirect(url_for('index'))
    

@app.route('/download_contrato/<path:filename>')
def download_contrato(filename):
    app.logger.info(f"Requisição de download para o arquivo: {filename}")
    # O diretório base para os contratos gerados (dentro da pasta de uploads)
    diretorio_contratos = os.path.join(app.config['UPLOAD_FOLDER'], 'contratos_gerados')
    app.logger.info(f"Tentando servir arquivo de: {diretorio_contratos}, arquivo: {filename}")
    try:
        return send_from_directory(directory=diretorio_contratos,
                                   path=filename, # Use 'path' em vez de 'filename' como argumento aqui
                                   as_attachment=True) # as_attachment=True força o download
    except FileNotFoundError:
        app.logger.error(f"Arquivo de contrato para download NÃO ENCONTRADO: {os.path.join(diretorio_contratos, filename)}")
        flash("Erro: Arquivo de contrato não encontrado para download.", "danger")
        return redirect(url_for('index'))
    except Exception as e:
        app.logger.error(f"Erro ao tentar enviar arquivo para download: {e}", exc_info=True)
        flash(f"Erro interno ao processar o download: {e}", "danger")
        return redirect(url_for('index'))
    
    
def start_server(host='0.0.0.0', port=8080): # Adicionei os parâmetros aqui
    # Removi o return daqui, app.run() bloqueia.
    app.run(host=host, port=port, debug=True)

if __name__ == '__main__':
    # Certifique-se que a pasta de uploads existe no início
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    start_server(port=8000) # Use a porta que preferir


def start_server(host='0.0.0.0', port=8080):
    return app.run(host=host, 
                   port=port, debug=True)

if __name__ == '__main__':
    start_server(port=8000)